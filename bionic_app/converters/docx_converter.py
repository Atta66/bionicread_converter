from __future__ import annotations

from io import BytesIO
from typing import Iterable

from bs4 import BeautifulSoup
from ebooklib import epub, ITEM_DOCUMENT
from docx import Document

from bionic_app.bionic_text import WORD_RE, bionic_prefix_len, flatten_text_blocks


def _add_bionic_paragraph(document: Document, text: str, ratio: float) -> None:
    paragraph = document.add_paragraph()

    for token in WORD_RE.findall(text):
        if token.isspace():
            paragraph.add_run(token)
            continue

        if not any(char.isalpha() for char in token):
            paragraph.add_run(token)
            continue

        prefix_len = bionic_prefix_len(token, ratio)
        strong_run = paragraph.add_run(token[:prefix_len])
        strong_run.bold = True
        if prefix_len < len(token):
            paragraph.add_run(token[prefix_len:])


def convert_pdf_text_to_docx(text_by_page: Iterable[str], ratio: float = 0.5, title: str = "Bionic PDF Export") -> bytes:
    document = Document()
    document.add_heading(title, level=1)

    pages = list(text_by_page)
    for index, page_text in enumerate(pages, start=1):
        document.add_heading(f"Page {index}", level=2)
        text = flatten_text_blocks([page_text])
        if text:
            for paragraph in text.split("\n\n"):
                _add_bionic_paragraph(document, paragraph, ratio)
        if index < len(pages):
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def convert_epub_to_docx(input_bytes: bytes, ratio: float = 0.5, title: str = "Bionic EPUB Export") -> bytes:
    book = epub.read_epub(BytesIO(input_bytes))
    document = Document()
    document.add_heading(title, level=1)

    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue

        soup = BeautifulSoup(item.get_content(), "lxml")
        text = soup.get_text("\n", strip=True)
        if not text:
            continue

        for paragraph in text.split("\n"):
            if paragraph.strip():
                _add_bionic_paragraph(document, paragraph.strip(), ratio)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
